# ===============================
# ðŸ“„ Analizador de PQR (Triple A) â€” Streamlit Cloud
# ===============================

import io
import os
import re
import zipfile
import base64
import json
import tempfile
from typing import List, Dict

import streamlit as st
import pandas as pd
from PIL import Image
import pypdfium2 as pdfium

# ======================
# ConfiguraciÃ³n Streamlit
# ======================
st.set_page_config(page_title="ðŸ“„ Analizador PQR â€“ Triple A", layout="wide")
st.title("ðŸ“„ Analizador de Derechos de PeticiÃ³n (PQR) â€“ Triple A ðŸ¦¾ðŸ¤–ðŸ«†")
st.caption("Carga PDFs o un ZIP con varios clientes. El sistema convierte a imÃ¡genes y extrae datos + pretensiones exactas.")

# ======================
# OpenAI Client (VisiÃ³n)
# ======================
try:
    from openai import OpenAI
    OPENAI_API_KEY = st.secrets.get("OPENAI_API_KEY", None)
    if OPENAI_API_KEY:
        os.environ["OPENAI_API_KEY"] = OPENAI_API_KEY
    oai_client = OpenAI()
    OAI_READY = True
except Exception as e:
    OAI_READY = False
    st.warning("âš ï¸ OpenAI SDK no disponible. Verifica requirements y secrets.")

# ======================
# Utilidades
# ======================
def pdf_bytes_to_images(pdf_bytes: bytes, dpi: int = 180) -> List[Image.Image]:
    """Convierte bytes de PDF a lista de PIL Images (todas las pÃ¡ginas)."""
    images: List[Image.Image] = []
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(pdf_bytes)
        tmp.flush()
        pdf = pdfium.PdfDocument(tmp.name)
        for i in range(len(pdf)):
            page = pdf[i]
            bitmap = page.render(scale=dpi/72).to_pil()
            images.append(bitmap.convert("RGB"))
    return images

def image_to_data_url(img: Image.Image, format: str = "JPEG", quality: int = 85) -> str:
    """Convierte PIL Image a data URL base64 (para enviar a VisiÃ³n)."""
    buf = io.BytesIO()
    img.save(buf, format=format, quality=quality)
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    return f"data:image/{format.lower()};base64,{b64}"

PROMPT_CORE = (
    "Eres analista de PQR de Triple A (Barranquilla). "
   "Analiza TODO el contenido visible del derecho de peticiÃ³n, incluyendo encabezados y secciones finales (como 'Notificaciones', 'DirecciÃ³n', 'Correo', 'C.C.', 'Atentamente'), "
    "pero ignora anexos, comprobantes o facturas. Si los datos del peticionario (nombre, cÃ©dula, correo, direcciÃ³n) aparecen en encabezado o firma, deben incluirse en la respuesta.\n\n""EXTRAE EXCLUSIVAMENTE si aparecen CLAROS y LITERALES estos campos (NO inventes nada):\n"
    "- CEDULA\n- NOMBRE COMPLETO\n- CORREO ELECTRÃ“NICO\n- TELÃ‰FONO\n- MOTIVO_PQR (puede estar en 'Referencia', 'Ref.', 'Asunto' o en la introducciÃ³n)\n"
    "- QUIEN_PRESENTA (peticionario principal, no el apoderado)\n- NOTIFICACION_A (direcciÃ³n/lugar de notificaciÃ³n; puede aparecer como 'Notificaciones', 'CitaciÃ³n', 'Correspondencia')\n"
    "- RESUMEN_PQR (sÃ­ntesis amplia y descriptiva del caso, de 8 a 12 lÃ­neas como mÃ­nimo, tomando frases del documento sin inventar. "
    "Debe incluir: el motivo, hechos mÃ¡s relevantes, reclamos formulados y contexto del servicio â€”por ejemplo, si es por facturaciÃ³n estimada, revisiÃ³n de medidor, cobros excesivos, suspensiÃ³n injusta o errores en la lecturaâ€”.)\n\n"
    "PRETENSIONES / SOLICITUDES / PETICIONES / EXIGENCIAS:\n"
    "- Captura TODAS las frases que expresen solicitudes, peticiones o exigencias del ciudadano.\n"
    "- Reconoce encabezados posibles: PETICIONES, PRETENSIONES, SOLICITUDES, EXIGENCIAS, REQUERIMIENTOS, DEMANDAS.\n"
    "- TambiÃ©n detecta cuando el texto usa verbos de acciÃ³n o ruego (aunque no haya tÃ­tulo): "
    "'Solicito', 'Pido', 'Requiero', 'Exijo', 'Se sirvan', 'Que se ordene', 'Mientras tanto', 'Ruego', 'Agradezco se sirvan', 'Solicita', 'Que se suspenda', 'Que se revoque', etc.\n"
    "- Si estÃ¡n enumeradas (1., 2., 3. o viÃ±etas), copia CADA Ã­tem literal como pretensiÃ³n independiente.\n"
    "- Si estÃ¡n dentro de pÃ¡rrafos corridos, extrae cada oraciÃ³n que contenga alguno de esos verbos.\n"
    "- NO resumas ni combines. Cada frase = 1 pretensiÃ³n.\n"
    "- Devuelve SIEMPRE: 'PRETENSIONES': {'TOTAL': N, 'DETALLE': ['1. ...', '2. ...', ...]}.\n\n"
    "REGLAS DE SALIDA:\n"
    "- Si un campo NO aparece o NO es claro, escribe exactamente: 'NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE'.\n"
    "- Devuelve SOLO un JSON vÃ¡lido con estas 9 claves fijas: \n"
    "  ['CEDULA','NOMBRE','CORREO','TELEFONO','MOTIVO_PQR','QUIEN_PRESENTA','NOTIFICACION_A','RESUMEN_PQR','PRETENSIONES']\n"
    "Recuerda que los datos de contacto (correo, telÃ©fono, direcciÃ³n) suelen estar cerca de las palabras 'Notificaciones', 'Correo', 'DirecciÃ³n' o 'TelÃ©fono'. "
    "Si los encuentras, extrÃ¡elos literalmente aunque estÃ©n fuera del pÃ¡rrafo principal.\n"

)

def call_vision_on_images(imgs: List[Image.Image], model: str = "gpt-4o-mini") -> Dict:
    """EnvÃ­a todas las pÃ¡ginas de un cliente al modelo de visiÃ³n con prompt blindado."""
    if not OAI_READY:
        return _manual_review_payload()

    content_blocks = [{"type": "text", "text": PROMPT_CORE}]
    for im in imgs:
        content_blocks.append({
            "type": "image_url",
            "image_url": {"url": image_to_data_url(im)}
        })

    try:
        resp = oai_client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": content_blocks}],
            temperature=0,
        )
        raw = resp.choices[0].message.content
        data = _safe_json_loads(raw)
        if not _valid_payload(data):
            return _manual_review_payload()
        return _normalize_payload(data)
    except Exception:
        return _manual_review_payload()

def _safe_json_loads(s: str):
    try:
        return json.loads(s)
    except Exception:
        s2 = s.strip()
        if s2.startswith("```") and s2.endswith("```"):
            s2 = "\n".join(s2.splitlines()[1:-1])
        try:
            return json.loads(s2)
        except Exception:
            return None

def _valid_payload(obj) -> bool:
    if not isinstance(obj, dict):
        return False
    keys = [
        "CEDULA","NOMBRE","CORREO","TELEFONO","MOTIVO_PQR",
        "QUIEN_PRESENTA","NOTIFICACION_A","RESUMEN_PQR","PRETENSIONES"
    ]
    for k in keys:
        if k not in obj:
            return False
    if not isinstance(obj.get("PRETENSIONES", {}), dict):
        return False
    return True

def _manual_review_payload() -> Dict:
    return {
        "CEDULA": "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE",
        "NOMBRE": "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE",
        "CORREO": "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE",
        "TELEFONO": "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE",
        "MOTIVO_PQR": "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE",
        "QUIEN_PRESENTA": "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE",
        "NOTIFICACION_A": "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE",
        "RESUMEN_PQR": "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE",
        "PRETENSIONES": {"TOTAL": 0, "DETALLE": ["NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"]},
    }

def _normalize_payload(d: Dict) -> Dict:
    def nz(v):
        v = (v or "").strip() if isinstance(v, str) else v
        return v if v else "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"

    pret = d.get("PRETENSIONES", {}) or {}
    detalle = pret.get("DETALLE", []) or []
    detalle_norm = []
    for i, item in enumerate(detalle, start=1):
        if isinstance(item, dict):
            item = item.get("texto") or item.get("detalle") or json.dumps(item, ensure_ascii=False)
        s = str(item).strip()
        if not s:
            s = "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"
        if not re.match(r"^\d+\.\s", s):
            s = f"{i}. {s}"
        detalle_norm.append(s)
    total = len(detalle_norm)

    return {
        "CEDULA": nz(d.get("CEDULA")),
        "NOMBRE": nz(d.get("NOMBRE")),
        "CORREO": nz(d.get("CORREO")),
        "TELEFONO": nz(d.get("TELEFONO")),
        "MOTIVO_PQR": nz(d.get("MOTIVO_PQR")),
        "QUIEN_PRESENTA": nz(d.get("QUIEN_PRESENTA")),
        "NOTIFICACION_A": nz(d.get("NOTIFICACION_A")),
        "RESUMEN_PQR": nz(d.get("RESUMEN_PQR")),
        "PRETENSIONES": {"TOTAL": total, "DETALLE": detalle_norm},
    }

# ======================
# Ingesta de archivos
# ======================
mode = st.radio("Modo de carga", ["Subir PDFs/ImÃ¡genes", "Subir ZIP"], horizontal=True)

client_docs: Dict[str, List[Image.Image]] = {}

if mode == "Subir ZIP":
    up = st.file_uploader("Sube un ZIP con varios PDFs/imagenes (cada PDF = un cliente)", type=["zip"])
    if up is not None:
        with zipfile.ZipFile(up) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                name = info.filename
                with zf.open(info) as f:
                    data = f.read()
                ext = os.path.splitext(name)[1].lower()
                base = os.path.splitext(os.path.basename(name))[0]
                if ext in [".pdf"]:
                    imgs = pdf_bytes_to_images(data)
                    client_docs[base] = client_docs.get(base, []) + imgs
                elif ext in [".jpg", ".jpeg", ".png"]:
                    img = Image.open(io.BytesIO(data)).convert("RGB")
                    group = base.split("_")[0]
                    client_docs[group] = client_docs.get(group, []) + [img]

else:
    ups = st.file_uploader(
        "Sube 1..N PDFs o imÃ¡genes (PNG/JPG)",
        type=["pdf", "png", "jpg", "jpeg"],
        accept_multiple_files=True
    )
    if ups:
        for up in ups:
            ext = os.path.splitext(up.name)[1].lower()
            base = os.path.splitext(os.path.basename(up.name))[0]
            if ext == ".pdf":
                imgs = pdf_bytes_to_images(up.read())
                client_docs[base] = client_docs.get(base, []) + imgs
            else:
                img = Image.open(up).convert("RGB")
                group = base.split("_")[0]
                client_docs[group] = client_docs.get(group, []) + [img]

# ======================
# Procesamiento
# ======================
if client_docs:
    st.subheader("ðŸ“¦ Clientes detectados")
    st.write(f"Total clientes: **{len(client_docs)}**")

    model_name = st.selectbox("Modelo de visiÃ³n", ["gpt-4o-mini", "gpt-4o"], index=0)
    run = st.button("ðŸš€ Procesar todo (extraer datos + pretensiones)")

    if run:
        rows = []
        qc_rows = []

        progress = st.progress(0)
        for idx, (client_key, pages_imgs) in enumerate(client_docs.items(), start=1):
            result = call_vision_on_images(pages_imgs, model=model_name)

            pret = result.get("PRETENSIONES", {})
            total = pret.get("TOTAL", 0)
            detalle_str = "\n".join(pret.get("DETALLE", []))

            row = {
                "CLIENTE_ID": client_key,
                "CEDULA": result.get("CEDULA"),
                "NOMBRE": result.get("NOMBRE"),
                "CORREO": result.get("CORREO"),
                "TELEFONO": result.get("TELEFONO"),
                "MOTIVO_PQR": result.get("MOTIVO_PQR"),
                "QUIEN_PRESENTA": result.get("QUIEN_PRESENTA"),
                "NOTIFICACION_A": result.get("NOTIFICACION_A"),
                "RESUMEN_PQR": result.get("RESUMEN_PQR"),
                "PRETENSIONES_TOTAL": total,
                "PRETENSIONES_DETALLE": detalle_str,
            }
            rows.append(row)

            qc_rows.append({
                "CLIENTE_ID": client_key,
                "PRETENSIONES_TOTAL": total,
                "ALERTA": "âš ï¸ REVISAR MANUALMENTE" if total == 0 else "OK",
            })

            progress.progress(min(idx/len(client_docs), 1.0))

        df = pd.DataFrame(rows)
        qc = pd.DataFrame(qc_rows)

        st.success("âœ”ï¸ Procesamiento completado.")
        st.subheader("ðŸ“Š Resultado consolidado")
        st.dataframe(df, use_container_width=True)

        st.subheader("ðŸ§ª Control de calidad (conteo de pretensiones)")
        st.dataframe(qc, use_container_width=True)

        out = io.BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="PQR", index=False)
            qc.to_excel(writer, sheet_name="QC_Pretensiones", index=False)
        st.download_button(
            "ðŸ“¥ Descargar Excel consolidado",
            data=out.getvalue(),
            file_name="pqr_consolidado.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
else:
    st.info("Sube PDFs, imÃ¡genes o un ZIP para comenzar.")
    # ======================
# ðŸ”— Emparejamiento de oficio (PDF) y respuesta (Word)
# ======================

st.divider()
st.header("ðŸ“‘ VALIDACIÃ“N DE RESPUESTAS A PQR â€” FASE 2")
st.caption("Sube los oficios en PDF y las respuestas en Word (DOCX). El sistema compararÃ¡ automÃ¡ticamente pretensiones y datos de notificaciÃ³n.")

uploaded_files = st.file_uploader(
    "Arrastra o selecciona 1 o varios archivos PDF (oficios) y DOCX (respuestas)",
    type=["pdf", "docx"],
    accept_multiple_files=True
)

pairs = {}

if uploaded_files:
    for up in uploaded_files:
        name = up.name
        base = os.path.splitext(name)[0]
        data = up.read()

        # Detectar nÃºmero de radicado o pÃ³liza (6+ dÃ­gitos)
        match = re.search(r"(\d{6,})", base)
        if not match:
            continue
        rad = match.group(1)

        ext = os.path.splitext(name)[1].lower()
        if ext == ".pdf":
            pairs.setdefault(rad, {})["pqr"] = data
            pairs[rad]["pqr_name"] = name
        elif ext == ".docx":
            pairs.setdefault(rad, {})["resp"] = data
            pairs[rad]["resp_name"] = name

    total_pairs = sum(1 for p in pairs.values() if "pqr" in p and "resp" in p)
    st.success(f"ðŸ“‚ Pares detectados: {total_pairs}")
    if total_pairs == 0:
        st.warning("âš ï¸ No se detectaron pares completos. Verifica que los nombres contengan el mismo nÃºmero de radicado o pÃ³liza.")


# ======================
# ðŸ“š FunciÃ³n comparativa IA (oficio vs respuesta)
# ======================
from docx import Document

def call_vision_on_pair(oficio_imgs: List[Image.Image], respuesta_bytes: bytes, model: str = "gpt-4o-mini") -> Dict:
    """Analiza un oficio (PDF en imÃ¡genes) y su respuesta (Word) usando GPT visiÃ³n+texto."""
    if not OAI_READY:
        return {"OBSERVACIONES": "IA no inicializada", "PRETENSIONES_TOTAL": 0}

    # 1ï¸âƒ£ Convertir oficio PDF â†’ base64
    oficio_blocks = []
    for im in oficio_imgs:
        oficio_blocks.append({
            "type": "image_url",
            "image_url": {"url": image_to_data_url(im)}
        })

    # 2ï¸âƒ£ Extraer texto de la respuesta Word
    try:
        doc = Document(io.BytesIO(respuesta_bytes))
        respuesta_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception:
        respuesta_text = "NO SE PUDO LEER LA RESPUESTA â€“ VALIDAR MANUALMENTE"

    # 3ï¸âƒ£ Prompt de contraste
    PROMPT_CONTRASTE = (
        "Eres un analista jurÃ­dico de PQR de Triple A en Barranquilla. "
        "Analiza el derecho de peticiÃ³n (oficio) y la respuesta (documento Word). "
        "Tu tarea es comparar ambos para determinar: \n\n"
        "1ï¸âƒ£ Si cada pretensiÃ³n o solicitud planteada por el ciudadano fue respondida completa o parcialmente. "
        "EvalÃºa por contenido, no por redacciÃ³n literal. Si se responde con otro texto pero satisface la solicitud, mÃ¡rcala como 'Respondida'. "
        "Si se omite o solo se menciona sin resolverla, mÃ¡rcala como 'No respondida'.\n\n"
        "2ï¸âƒ£ Verifica si los datos del peticionario (nombre, cÃ©dula, correo, direcciÃ³n) son coherentes entre el oficio y la respuesta. "
        "Considera que pequeÃ±as diferencias tipogrÃ¡ficas (mayÃºsculas/minÃºsculas, confusiÃ³n entre 'l', '1' o 'i', espacios o acentos) NO constituyen error. "
        "Solo marca como 'Incorrectos' si hay cambio de persona o correo completamente distinto. "
        "Si hay duda leve, marca 'SÃ­ (diferencia menor)' y explica.\n\n"
        "3ï¸âƒ£ Si hay errores de digitaciÃ³n, menciona ejemplos especÃ­ficos (por ejemplo, correo con un carÃ¡cter cambiado, omisiÃ³n de nÃºmero de cÃ©dula, etc.).\n\n"
        "Devuelve un JSON con las claves fijas: \n"
        "['NOMBRE','CEDULA','CORREO','NOTIFICACION_A','PRETENSIONES_TOTAL','PRETENSIONES_DETALLE',"
        "'PRETENSIONES_CORRECTAS','DATOS_NOTIFICACION_CORRECTOS','OBSERVACIONES']\n"
        "Donde 'PRETENSIONES_DETALLE' es una lista de las pretensiones numeradas con indicador (Respondida / No respondida)."
    )

    # 4ï¸âƒ£ Enviar a OpenAI
    content_blocks = [{"type": "text", "text": PROMPT_CONTRASTE}]
    content_blocks.extend(oficio_blocks)
    content_blocks.append({"type": "text", "text": f"Respuesta del analista:\n{respuesta_text}"})

    try:
        resp = oai_client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": content_blocks}],
            temperature=0,
        )
        raw = resp.choices[0].message.content
        data = _safe_json_loads(raw)
        if not data or not isinstance(data, dict):
            raise ValueError("Respuesta IA invÃ¡lida o vacÃ­a")
        return data
    except Exception as e:
        return {"OBSERVACIONES": f"Falla en procesamiento IA: {e}", "PRETENSIONES_TOTAL": 0}


# ======================
# âš™ï¸ Procesamiento principal â€“ ComparaciÃ³n masiva
# ======================
if pairs:
    st.subheader("âš™ï¸ Procesamiento de anÃ¡lisis comparativo")
    model_name = st.selectbox("Modelo de visiÃ³n", ["gpt-4o-mini", "gpt-4o"], index=0)
    run = st.button("ðŸš€ Analizar todos los pares")

    if run:
        resultados = []
        progreso = st.progress(0)

        for i, (rad, docs) in enumerate(pairs.items(), start=1):
            try:
                oficio_imgs = pdf_bytes_to_images(docs["pqr"])
                respuesta_bytes = docs["resp"]

                resultado = call_vision_on_pair(oficio_imgs, respuesta_bytes, model=model_name)
                resultado["RADICADO"] = rad
                resultado["ARCHIVO_PQR"] = docs.get("pqr_name", "NO REGISTRADO")
                resultado["ARCHIVO_RESPUESTA"] = docs.get("resp_name", "NO REGISTRADO")

                # ======================
                # ðŸ”§ NormalizaciÃ³n y mÃ©tricas
                # ======================
                detalle = resultado.get("PRETENSIONES_DETALLE", [])
                respondidas = 0
                total_pretensiones = 0

                if isinstance(detalle, list) and detalle:
                    detalle_limpio = []
                    for idx, p in enumerate(detalle):
                        if isinstance(p, dict):
                            texto = p.get("pretension") or p.get("texto") or str(p)
                            estado = p.get("estado") or p.get("respondida") or "NO"
                        else:
                            texto = str(p)
                            estado = "NO"

                        icono = "âœ…" if estado.lower().strip() in ["respondida", "sÃ­", "si", "completa", "totalmente respondida"] else "âŒ"
                        detalle_limpio.append(f"{idx+1}. {texto.strip()} ({icono})")

                        total_pretensiones += 1
                        if icono == "âœ…":
                            respondidas += 1

                    detalle_str = "\n".join(detalle_limpio)
                else:
                    detalle_str = "NO SE ENCONTRARON PRETENSIONES â€“ VALIDAR MANUALMENTE"
                    total_pretensiones = 0
                    respondidas = 0

                porcentaje_respondidas = (respondidas / total_pretensiones * 100) if total_pretensiones > 0 else 0

                # Agregar fila consolidada
                resultados.append({
                    "RADICADO": rad,
                    "POLIZA": re.search(r'(\d{6,})', docs.get("pqr_name", "")) and re.search(r'(\d{6,})', docs["pqr_name"]).group(1) or "NO SE APORTÃ“",
                    "NOMBRE": resultado.get("NOMBRE", "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"),
                    "CEDULA": resultado.get("CEDULA", "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"),
                    "CORREO": resultado.get("CORREO", "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"),
                    "NOTIFICACION_A": resultado.get("NOTIFICACION_A", "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"),
                    "PRETENSIONES_TOTAL": total_pretensiones,
                    "PRETENSIONES_RESPONDIDAS": respondidas,
                    "PORCENTAJE_RESPONDIDAS": round(porcentaje_respondidas, 1),
                    "PRETENSIONES_DETALLE": detalle_str,
                    "PRETENSIONES_CORRECTAS": resultado.get("PRETENSIONES_CORRECTAS", "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"),
                    "DATOS_NOTIFICACION_CORRECTOS": resultado.get("DATOS_NOTIFICACION_CORRECTOS", "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"),
                    "OBSERVACIONES": resultado.get("OBSERVACIONES", "NO SE APORTÃ“ â€“ VALIDAR MANUALMENTE"),
                    "ARCHIVO_PQR": docs.get("pqr_name", "NO REGISTRADO"),
                    "ARCHIVO_RESPUESTA": docs.get("resp_name", "NO REGISTRADO"),
                })

            except Exception as e:
                resultados.append({
                    "RADICADO": rad,
                    "POLIZA": "ERROR",
                    "NOMBRE": "ERROR",
                    "CEDULA": "ERROR",
                    "CORREO": "ERROR",
                    "NOTIFICACION_A": "ERROR",
                    "PRETENSIONES_TOTAL": 0,
                    "PRETENSIONES_RESPONDIDAS": 0,
                    "PORCENTAJE_RESPONDIDAS": 0,
                    "PRETENSIONES_DETALLE": str(e),
                    "PRETENSIONES_CORRECTAS": "ERROR",
                    "DATOS_NOTIFICACION_CORRECTOS": "ERROR",
                    "OBSERVACIONES": "Falla en procesamiento IA",
                    "ARCHIVO_PQR": docs.get("pqr_name", "NO REGISTRADO"),
                    "ARCHIVO_RESPUESTA": docs.get("resp_name", "NO REGISTRADO"),
                })

            progreso.progress(i / len(pairs))

        # ðŸ§¾ Crear DataFrame
        df = pd.DataFrame(resultados)
        st.success("âœ… AnÃ¡lisis completado con Ã©xito.")
        st.subheader("ðŸ“Š Resultado consolidado")
        st.dataframe(df, use_container_width=True)

        # ðŸ“¥ Generar Excel con formato condicional
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="Comparativo_PQR", index=False)
            ws = writer.sheets["Comparativo_PQR"]

            from openpyxl.formatting.rule import CellIsRule
            from openpyxl.styles import PatternFill

            fill_green = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
            fill_yellow = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
            fill_red = PatternFill(start_color="F8CBAD", end_color="F8CBAD", fill_type="solid")

            # âœ… Colores por % respondidas
            ws.conditional_formatting.add("K2:K2000", CellIsRule(operator="greaterThanOrEqual", formula=["80"], fill=fill_green))
            ws.conditional_formatting.add("K2:K2000", CellIsRule(operator="between", formula=["40", "79.9"], fill=fill_yellow))
            ws.conditional_formatting.add("K2:K2000", CellIsRule(operator="lessThan", formula=["40"], fill=fill_red))

            # ðŸŸ© Datos correctos
            ws.conditional_formatting.add("I2:I2000", CellIsRule(operator="equal", formula=['"SÃ­"'], fill=fill_green))
            ws.conditional_formatting.add("I2:I2000", CellIsRule(operator="equal", formula=['"No"'], fill=fill_red))

        st.download_button(
            "ðŸ“¥ Descargar Excel consolidado con formato",
            data=out.getvalue(),
            file_name="comparativo_pqr_respuestas.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

else:
    st.info("Sube los documentos para analizar (PDF + DOCX con mismo radicado).")
